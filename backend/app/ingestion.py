"""Turn what a user gives us into plain text chunks with provenance.

Everything an agent later says about a person has to be traceable to one of these
chunks. That is why each chunk keeps its source and an ordinal — a claim without a
retrievable chunk behind it is a hallucination, and the persona layer is expected to
say "not in profile" rather than invent one.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup

SUPPORTED_UPLOAD_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".markdown"}

STRIPPED_HTML_TAGS = (
    "script",
    "style",
    "noscript",
    "nav",
    "footer",
    "header",
    "form",
    "svg",
    "iframe",
)


class IngestionError(ValueError):
    """A source could not be read. The message is safe to show a user."""


@dataclass(frozen=True)
class ExtractedSource:
    title: str
    text: str
    kind: str
    detail: str


def _collapse_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()


def extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise IngestionError("That PDF could not be read. Try exporting it again.") from exc
    return _collapse_whitespace("\n\n".join(pages))


def extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text for cell in row.cells))
    except Exception as exc:
        raise IngestionError("That Word document could not be read.") from exc
    return _collapse_whitespace("\n".join(blocks))


def extract_plain_text(data: bytes) -> str:
    try:
        return _collapse_whitespace(data.decode("utf-8"))
    except UnicodeDecodeError:
        return _collapse_whitespace(data.decode("utf-8", errors="replace"))


def extract_upload(filename: str, data: bytes, *, max_characters: int) -> ExtractedSource:
    suffix = ""
    if "." in filename:
        suffix = filename[filename.rfind(".") :].lower()
    if suffix not in SUPPORTED_UPLOAD_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_UPLOAD_SUFFIXES))
        raise IngestionError(
            f"Unsupported file type '{suffix or filename}'. Supported: {supported}"
        )

    if suffix == ".pdf":
        text = extract_pdf(data)
    elif suffix == ".docx":
        text = extract_docx(data)
    else:
        text = extract_plain_text(data)

    if not text.strip():
        raise IngestionError(
            "No readable text found in that file. Scanned images are not supported yet."
        )
    return ExtractedSource(
        title=filename,
        text=text[:max_characters],
        kind="upload",
        detail=filename,
    )


def normalize_url(raw_url: str) -> str:
    url = raw_url.strip()
    if not url:
        raise IngestionError("Enter a URL first.")
    if not url.startswith(("http://", "https://")):
        url = f"https://{url}"
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise IngestionError("That does not look like a valid web address.")
    return url


def extract_html(html: str) -> tuple[str, str]:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(list(STRIPPED_HTML_TAGS)):
        tag.decompose()
    title = ""
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    if not title:
        heading = soup.find("h1")
        if heading:
            title = heading.get_text(" ", strip=True)
    return title, _collapse_whitespace(soup.get_text("\n", strip=True))


def fetch_url(raw_url: str, *, timeout: float, max_characters: int) -> ExtractedSource:
    url = normalize_url(raw_url)
    try:
        response = httpx.get(
            url,
            timeout=timeout,
            follow_redirects=True,
            headers={"User-Agent": "AgentCircleBot/0.1 (+persona ingestion)"},
        )
        response.raise_for_status()
    except httpx.HTTPError as exc:
        raise IngestionError(f"Could not fetch {url}. It may be private or offline.") from exc

    content_type = response.headers.get("content-type", "")
    if "application/pdf" in content_type:
        text = extract_pdf(response.content)
        title = url
    elif "html" in content_type or not content_type:
        title, text = extract_html(response.text)
    else:
        text = _collapse_whitespace(response.text)
        title = url

    if not text.strip():
        raise IngestionError(f"No readable text found at {url}.")
    return ExtractedSource(
        title=title or url,
        text=text[:max_characters],
        kind="link",
        detail=url,
    )


def chunk_text(text: str, *, size: int, overlap: int) -> list[str]:
    """Split on paragraph boundaries, packing up to `size` characters per chunk.

    Overlap carries the tail of the previous chunk forward so a fact spanning a
    boundary stays retrievable from at least one chunk.
    """
    cleaned = _collapse_whitespace(text)
    if not cleaned:
        return []
    if size <= 0:
        return [cleaned]
    overlap = max(0, min(overlap, size // 2))

    paragraphs = [block.strip() for block in cleaned.split("\n\n") if block.strip()]
    chunks: list[str] = []
    current = ""

    def flush() -> None:
        nonlocal current
        if current.strip():
            chunks.append(current.strip())
        current = ""

    for paragraph in paragraphs:
        # A single oversized paragraph is hard-split rather than dropped.
        while len(paragraph) > size:
            flush()
            chunks.append(paragraph[:size].strip())
            paragraph = paragraph[size - overlap :]
        if not current:
            current = paragraph
        elif len(current) + len(paragraph) + 2 <= size:
            current = f"{current}\n\n{paragraph}"
        else:
            tail = current[-overlap:] if overlap else ""
            flush()
            current = f"{tail}\n\n{paragraph}".strip() if tail else paragraph
    flush()
    return [chunk for chunk in chunks if chunk]
